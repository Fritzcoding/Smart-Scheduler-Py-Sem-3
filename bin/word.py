
# Export and visualization helpers:
# - `docx` + `docx.shared.Pt`: build .docx schedule documents from a template
# - `deepcopy`: used to duplicate document table structures when creating group tables
# - `numpy` and `matplotlib`: simple numeric aggregation and visualization (bar charts)
# - `PIL.Image*` and `cv2`: create and manipulate schedule images (Pillow preferred for simple drawing)
# - `json`/`csv`/`os`: write schedule exports and manage filesystem
import docx
from copy import deepcopy
from docx.shared import Pt 
import numpy as np
import matplotlib.pyplot as plt
from PIL import Image, ImageDraw, ImageFont
import cv2
import json
import csv
import os


GAMES = ["name games","softball", "basketball", "squash", "ultimate", "hockey", "lacrosse", "football", "tennis", "volleyball", "soccer"]

# maximum number of periods per group
MAX_PERIODS = 6


def cap_periods(matrix, max_periods=MAX_PERIODS):
	"""
	Ensure no group has more than `max_periods` periods.
	Trims extra periods while preserving original order.
	"""
	capped = []
	for group in matrix:
		# if group is shorter than max_periods, keep as-is
		capped.append(group[:max_periods])
	return capped

def create_tables(n, doc):
	"""
	will create all of the neccessary tables/schdules for
	each week in a new word document.

	:param doc: word document to add tables to
	:param n: number of tables to add (for groups 2 onwards; group 1 is the template)
	:return: None
	"""
	for x in range(n):
		p = doc.add_paragraph("Group " + str(x+2))
		p.style = "group"
		template = doc.tables[0]
		tbl = template._tbl
		new_tbl = deepcopy(tbl)
		paragraph = doc.add_paragraph()
		paragraph._p.addnext(new_tbl)
		# add a page break after this new group unless it's the last one we add
		if x < (n - 1):
			doc.add_page_break()


def fill_tables(matrix, doc):
	"""
	fills in the tables previously created by the create_tables function:
	
	:param matrix: a 3d list
	:param doc: the word document
	:return: None
	"""
	for group, line in enumerate(doc.tables):
		r = 0
		for row in line.rows:
			col = 0
			update = False
			for cell in row.cells:
				if cell.text == "" or cell.text.lower().strip() == "name games":
					try:
						cell.text = matrix[group][r][col]
						if cell.text == "Name Games":
							cell.paragraphs[0].style = 'name'
						else:
							cell.paragraphs[0].style = 'center'
						col += 1
						update = True
					except:
						break
			if update:
				r += 1


def make_word_doc(matrix, file_name="Week 1"):
	"""
	creates the word document that contains the new schdule.

	:param matrix: 3d list
	:param file_name: str
	:return: None
	"""
	# enforce maximum periods per group
	matrix = cap_periods(matrix)

	doc = docx.Document("2019 Template Schedules.docx")
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Arial'
	font.size = Pt(12)
	groups = len(matrix)
	
	# The template should already contain the header for Group 1; do not override it.
	
	create_tables(groups-1, doc)
	os.makedirs("Generated Schedules", exist_ok=True)
	doc.save("Generated Schedules/" + file_name + " Schedules.docx")
	doc = docx.Document("Generated Schedules/" + file_name + " Schedules.docx")
	fill_tables(matrix, doc)
	doc.save("Generated Schedules/" + file_name + " Schedules.docx")
    
	# Generate additional outputs with new libraries
	export_schedule_json(matrix, file_name)
	export_schedule_csv(matrix, file_name)
	create_schedule_visualization(matrix, file_name)
	generate_schedule_image_pillow(matrix, file_name)


def export_schedule_json(matrix, file_name):
	"""
	Export schedule to JSON format using the json module.
	
	:param matrix: 3d list
	:param file_name: str
	:return: None
	"""
	try:
		# cap periods to MAX_PERIODS before exporting
		matrix = cap_periods(matrix)
		schedule_dict = {}
		for g, group in enumerate(matrix):
			schedule_dict[f"Group {g+1}"] = group
		
		filepath = f"Generated Schedules/{file_name}_schedule.json"
		with open(filepath, 'w') as f:
			json.dump(schedule_dict, f, indent=2)
		print(f"✓ JSON export successful: {filepath}")
	except Exception as e:
		print(f"✗ Error exporting JSON: {e}")


def export_schedule_csv(matrix, file_name):
	"""
	Export schedule to CSV format using the csv module.
	
	:param matrix: 3d list
	:param file_name: str
	:return: None
	"""
	try:
		# cap periods to MAX_PERIODS before exporting
		matrix = cap_periods(matrix)
		filepath = f"Generated Schedules/{file_name}_schedule.csv"
		with open(filepath, 'w', newline='', encoding='utf-8') as f:
			writer = csv.writer(f)
			writer.writerow(["Group", "Period", "Monday", "Tuesday", "Wednesday", "Thursday"])
            
			for g, group in enumerate(matrix):
				for period_num, period in enumerate(group):
					row = [f"Group {g+1}", f"Period {period_num+1}"] + period
					writer.writerow(row)
		print(f"✓ CSV export successful: {filepath}")
	except Exception as e:
		print(f"✗ Error exporting CSV: {e}")


def create_schedule_visualization(matrix, file_name):
	"""
	Create a visualization of schedule using Matplotlib.
	
	:param matrix: 3d list
	:param file_name: str
	:return: None
	"""
	try:
		# cap periods to MAX_PERIODS before visualizing
		matrix = cap_periods(matrix)
		# Count activities
		activity_counts = {}
		for group in matrix:
			for period in group:
				for activity in period:
					if activity and activity != "":
						activity_counts[activity] = activity_counts.get(activity, 0) + 1
		
		# Create visualization using matplotlib
		plt.figure(figsize=(12, 6))
		activities = list(activity_counts.keys())
		counts = list(activity_counts.values())
		
		plt.bar(activities, counts, color='steelblue')
		plt.xlabel('Activities')
		plt.ylabel('Frequency')
		plt.title(f'Schedule Activity Frequency - {file_name}')
		plt.xticks(rotation=45, ha='right')
		plt.tight_layout()
		
		# Save the figure
		filepath = f"Generated Schedules/{file_name}_visualization.png"
		plt.savefig(filepath, dpi=100)
		plt.close()
		print(f"✓ Matplotlib visualization created: {filepath}")
	except Exception as e:
		print(f"✗ Error creating visualization: {e}")


def generate_schedule_image_pillow(matrix, file_name):
	"""
	Generate schedule image using Pillow.
	
	:param matrix: 3d list
	:param file_name: str
	:return: None
	"""
	try:
		# cap periods to MAX_PERIODS before generating image
		matrix = cap_periods(matrix)
		# Create image
		img_width = 1400
		img_height = 200 + (len(matrix) * 180)
		img = Image.new('RGB', (img_width, img_height), color='white')
		draw = ImageDraw.Draw(img)
		
		# Add title
		draw.text((50, 20), f"Schedule - {file_name}", fill='black')
		
		# Add schedule data
		y_offset = 80
		for g, group in enumerate(matrix):
			draw.text((50, y_offset), f"Group {g+1}:", fill='darkblue')
			y_offset += 35
			
			for period_num, period in enumerate(group):
				period_text = f"Period {period_num+1}: {' | '.join(period)}"
				draw.text((70, y_offset), period_text, fill='black')
				y_offset += 25
			
			y_offset += 25
		
		# Save image
		filepath = f"Generated Schedules/{file_name}_schedule_pillow.png"
		img.save(filepath)
		print(f"✓ Pillow image generated: {filepath}")
	except Exception as e:
		print(f"✗ Error generating image: {e}")


def make_word_doc_only(matrix, file_name="Week 1"):
	"""
	Create and save only the Word document for the schedule (no JSON/CSV/images).

	:param matrix: 3d list
	:param file_name: str
	:return: None
	"""
	try:
		matrix = cap_periods(matrix)
		doc = docx.Document("2019 Template Schedules.docx")
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Arial'
		font.size = Pt(12)
		groups = len(matrix)

		create_tables(groups-1, doc)
		os.makedirs("Generated Schedules", exist_ok=True)
		doc.save("Generated Schedules/" + file_name + " Schedules.docx")
		doc = docx.Document("Generated Schedules/" + file_name + " Schedules.docx")
		fill_tables(matrix, doc)
		doc.save("Generated Schedules/" + file_name + " Schedules.docx")
		print(f"✓ Word export successful: Generated Schedules/{file_name} Schedules.docx")
	except Exception as e:
		print(f"✗ Error exporting Word-only document: {e}")


